# -*- coding: utf-8 -*-

# --- Документація ---
# Професійний парсер для сайту Swappa.com.
# ВЕРСІЯ 5.0: Гнучка конфігурація, стійкий до змін парсинг, підтримка 4 форматів.
#
# - Налаштування через аргументи командного рядка (URL, API-ключ, ім'я файлу).
# - Використовує карту колонок для надійного вилучення даних.
# - Зберігає звіти в форматах CSV, DOCX, JSON та XLSX.
#
# Автор: Ярослав (при содействии Gemini)

# --- 1. Імпорт необхідних бібліотек ---
import requests
import pandas as pd
from bs4 import BeautifulSoup
from docx import Document
import logging
import argparse  # Для обробки аргументів командного рядка
import json      # Для збереження в JSON

# --- 2. Налаштування логування ---
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# --- 3. Карта колонок для надійного парсингу ---
# Описує, в якій за рахунком колонці (індекс) знаходяться потрібні дані.
# Якщо Swappa змінить порядок колонок, достатньо буде змінити цифри тут.
COLUMN_MAPPING = {
    "price": 1,
    "carrier": 3,
    "color": 4,
    "storage": 5,
    "model": 6,
    "condition": 7,
    "battery": 8,
    "seller": 9,
    "location": 10,
    "shipping": 12,
    "code": 13
}


def fetch_page_html(api_key: str, url: str) -> str:
    """
    Відправляє запит до ScrapingBee API для отримання HTML-кода сторінки.
    """
    logging.info(f"Відправка запиту на отримання HTML через ScrapingBee для URL: {url}")
    response = requests.get(
        url='https://app.scrapingbee.com/api/v1/',
        params={'api_key': api_key, 'url': url, 'render_js': 'true'}
    )
    if response.status_code == 200:
        logging.info("HTML-код сторінки успішно отримано.")
        return response.text
    else:
        logging.error(f"Помилка при отриманні сторінки. Статус: {response.status_code}, Відповідь: {response.text}")
        return None


def parse_html_data(html_content: str) -> list:
    """
    Приймає HTML-код, розбирає його і вилучає дані з таблиці,
    використовуючи гнучку карту колонок COLUMN_MAPPING.
    """
    if not html_content:
        return []

    logging.info("Починаємо розбір (парсинг) HTML-коду...")
    soup = BeautifulSoup(html_content, 'html.parser')
    
    listing_elements = soup.select("table#listings_table tbody tr")
    logging.info(f"Знайдено оголошень на сторінці: {len(listing_elements)}")
    
    if not listing_elements:
        return []

    all_phones_data = []
    for i, row in enumerate(listing_elements, 1):
        try:
            cells = row.select("td")
            
            # Перевірка, що в рядку достатньо колонок для безпечного парсингу
            if len(cells) < max(COLUMN_MAPPING.values()) + 1:
                continue

            phone_data = {}
            for field_name, cell_index in COLUMN_MAPPING.items():
                # Спеціальна обробка для продавця, щоб взяти лише ім'я
                if field_name == "seller":
                    phone_data[field_name] = cells[cell_index].text.strip().split('\n')[0]
                else:
                    phone_data[field_name] = cells[cell_index].text.strip()
            
            all_phones_data.append(phone_data)
        except Exception as e:
            logging.warning(f"Не вдалося обробити рядок #{i}. Помилка: {e}. Пропускаємо.")
            continue
            
    return all_phones_data


# --- Функції для збереження даних у різних форматах ---

def save_to_csv(data: list, filename: str):
    if not data: return
    logging.info(f"Збереження {len(data)} записів у CSV: {filename}...")
    try:
        pd.DataFrame(data).to_csv(filename, index=False, encoding='utf-8')
        logging.info(f"Файл {filename} успішно збережено.")
    except Exception as e:
        logging.error(f"Помилка при збереженні в CSV: {e}")

def save_to_xlsx(data: list, filename: str):
    if not data: return
    logging.info(f"Збереження {len(data)} записів у XLSX: {filename}...")
    try:
        pd.DataFrame(data).to_excel(filename, index=False, engine='openpyxl')
        logging.info(f"Файл {filename} успішно збережено.")
    except Exception as e:
        logging.error(f"Помилка при збереженні в XLSX: {e}")

def save_to_json(data: list, filename: str):
    if not data: return
    logging.info(f"Збереження {len(data)} записів у JSON: {filename}...")
    try:
        with open(filename, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=4)
        logging.info(f"Файл {filename} успішно збережено.")
    except Exception as e:
        logging.error(f"Помилка при збереженні в JSON: {e}")

def save_to_docx(data: list, filename: str):
    if not data: return
    logging.info(f"Створення Word-звіту: {filename}...")
    try:
        document = Document()
        document.add_heading('Звіт по оголошенням Swappa', level=1)
        document.add_paragraph(f"Зібрано {len(data)} оголошень.")
        
        for item in data:
            title = f"{item.get('storage')} {item.get('color')} ({item.get('condition')})"
            document.add_heading(title, level=3)
            p = document.add_paragraph()
            p.add_run('Ціна: ').bold = True
            p.add_run(f"{item.get('price', 'N/A')} | ")
            p.add_run('Оператор: ').bold = True
            p.add_run(f"{item.get('carrier', 'N/A')}\n")
            p.add_run('Продавець: ').bold = True
            p.add_run(f"{item.get('seller', 'N/A')} ({item.get('location', 'N/A')})")

        document.save(filename)
        logging.info(f"Звіт {filename} успішно збережено.")
    except Exception as e:
        logging.error(f"Помилка при збереженні в DOCX: {e}")


def main():
    """Головна функція, що керує всім процесом."""
    parser = argparse.ArgumentParser(description="Парсер оголошень з сайту Swappa.com (v5.0)")
    parser.add_argument('--api_key', type=str, required=True, help='Ваш API ключ від ScrapingBee')
    parser.add_argument('--url', type=str, default="https://swappa.com/listings/apple-iphone-13-pro-max", help='URL сторінки для парсингу')
    parser.add_argument('--output', type=str, default="swappa_report", help="Базове ім'я для вихідних файлів (без розширення)")
    args = parser.parse_args()

    logging.info("Запуск парсера v5.0 (Flexible CLI Version)...")
    
    html_content = fetch_page_html(args.api_key, args.url)
    
    if html_content:
        scraped_data = parse_html_data(html_content)
        if scraped_data:
            # Створюємо імена файлів на основі аргументу --output
            base_filename = args.output
            # Зберігаємо дані в усіх форматах
            save_to_csv(scraped_data, f"{base_filename}.csv")
            save_to_xlsx(scraped_data, f"{base_filename}.xlsx")
            save_to_json(scraped_data, f"{base_filename}.json")
            save_to_docx(scraped_data, f"{base_filename}.docx")
        else:
            logging.warning("HTML отримано, але не вдалося вилучити дані. Перевірте CSS-селектори та структуру сторінки.")
            # Зберігаємо HTML для аналізу в разі помилки
            with open("debug_page.html", "w", encoding="utf-8") as f:
                f.write(html_content)
            logging.info("HTML-код сторінки збережено в debug_page.html для аналізу.")
    
    logging.info("Робота завершена.")


if __name__ == "__main__":
    main()