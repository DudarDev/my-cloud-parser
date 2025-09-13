# -*- coding: utf-8 -*-

# --- Документация ---
# Финальная версия парсера для сайта Swappa.com.
# ВЕРСИЯ 4.0: Адаптирован для парсинга табличного вида страницы.
# Использует API сервиса ScrapingBee для обхода Cloudflare.
#
# Автор: Ярослав (при содействии Gemini)

# --- 1. Импорт необходимых библиотек ---
import requests
import pandas as pd
from bs4 import BeautifulSoup
from docx import Document
import logging

# --- 2. Настройка логирования ---
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# --- 3. Конфигурация ---
SCRAPINGBEE_API_KEY = "9E3N7FKEJZZBQZ72RPSO7WF6DXO2XN6TM4XXZH3O2WS0T6ZYYV370BIZB1R20KPWT0FTHECHSARCDET7"
TARGET_URL = "https://swappa.com/listings/apple-iphone-13-pro-max"
OUTPUT_CSV_FILE = "swappa_iphone_report.csv"
OUTPUT_DOCX_FILE = "swappa_iphone_report.docx"


def fetch_page_html(api_key: str, url: str) -> str:
    """
    Отправляет запрос к ScrapingBee API для получения HTML-кода страницы.
    """
    logging.info(f"Отправка запроса на получение HTML через ScrapingBee для URL: {url}")
    
    response = requests.get(
        url='https://app.scrapingbee.com/api/v1/',
        params={
            'api_key': api_key,
            'url': url,
            'render_js': 'true', # Говорим сервису выполнить JavaScript на странице
        }
    )
    
    if response.status_code == 200:
        logging.info("HTML-код страницы успешно получен.")
        return response.text
    else:
        logging.error(f"Ошибка при получении страницы. Статус: {response.status_code}, Ответ: {response.text}")
        return None


def parse_html_data(html_content: str) -> list:
    """
    Принимает HTML-код, разбирает его с помощью BeautifulSoup и извлекает данные
    из ТАБЛИЧНОГО представления.
    """
    if not html_content:
        return []

    logging.info("Начинаем разбор (парсинг) HTML-кода...")
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # 🔥 НОВЫЙ СЕЛЕКТОР для строк таблицы
    listing_elements = soup.select("table#listings_table tbody tr")
    logging.info(f"Найдено объявлений на странице: {len(listing_elements)}")
    
    if not listing_elements:
        return []

    all_phones_data = []
    for i, row in enumerate(listing_elements, 1):
        try:
            # 🔥 НОВЫЙ СПОСОБ ПОИСКА: ищем все ячейки (<td>) в строке
            cells = row.select("td")
            
            # Проверяем, что в строке достаточно ячеек, чтобы избежать ошибок
            if len(cells) < 14: # В таблице 14+ колонок
                continue

            # Извлекаем данные по точному порядку ячеек
            price = cells[1].text.strip()
            carrier = cells[3].text.strip()
            color = cells[4].text.strip()
            storage = cells[5].text.strip()
            model = cells[6].text.strip()
            condition = cells[7].text.strip()
            battery = cells[8].text.strip()
            seller = cells[9].text.strip().split('\n')[0] # Берем только имя, без рейтинга
            location = cells[10].text.strip()
            shipping = cells[12].text.strip()
            code = cells[13].text.strip()

            phone_data = {
                "price": price,
                "carrier": carrier,
                "color": color,
                "storage": storage,
                "model": model,
                "condition": condition,
                "battery": battery,
                "seller": seller,
                "location": location,
                "shipping": shipping,
                "code": code
            }
            all_phones_data.append(phone_data)
        except Exception as e:
            logging.warning(f"Не удалось обработать строку #{i}. Ошибка: {e}. Пропускаем.")
            continue
            
    return all_phones_data


def save_to_csv(data: list, filename: str):
    """Сохраняет список словарей в CSV-файл с помощью Pandas."""
    if not data:
        logging.warning("Нет данных для сохранения в CSV.")
        return
    logging.info(f"Сохранение {len(data)} записей в файл {filename}...")
    try:
        df = pd.DataFrame(data)
        df.to_csv(filename, index=False, encoding='utf-8')
        logging.info(f"Данные успешно сохранены в {filename}")
    except Exception as e:
        logging.error(f"Ошибка при сохранении в CSV: {e}")


def save_to_docx(data: list, filename: str):
    """Создает и сохраняет отчет в формате Word (.docx)."""
    if not data:
        logging.warning("Нет данных для сохранения в DOCX.")
        return
    logging.info(f"Создание Word-отчета и сохранение в {filename}...")
    try:
        document = Document()
        document.add_heading('Отчет по объявлениям Swappa: iPhone 13 Pro Max', level=1)
        document.add_paragraph(f"Собрано {len(data)} объявлений.")
        
        for item in data:
            title = f"{item.get('storage')} {item.get('color')} ({item.get('condition')})"
            document.add_heading(title, level=3)
            
            p = document.add_paragraph()
            p.add_run('Цена: ').bold = True
            p.add_run(item.get('price', 'N/A'))
            
            p.add_run(' | Оператор: ').bold = True
            p.add_run(item.get('carrier', 'N/A'))
            
            p.add_run('\nПродавец: ').bold = True
            p.add_run(f"{item.get('seller', 'N/A')} ({item.get('location', 'N/A')})")

        document.save(filename)
        logging.info(f"Отчет успешно сохранен в {filename}")
    except Exception as e:
        logging.error(f"Ошибка при сохранении в DOCX: {e}")


def main():
    """Главная функция, управляющая всем процессом."""
    logging.info("Запуск парсера v4.0 (API, Table Version)...")
    
    if "ВАШ_API_КЛЮЧ" in SCRAPINGBEE_API_KEY:
        logging.error("Пожалуйста, вставьте ваш API-ключ от ScrapingBee в переменную SCRAPINGBEE_API_KEY.")
        return

    html_content = fetch_page_html(SCRAPINGBEE_API_KEY, TARGET_URL)
    
    if html_content:
        scraped_data = parse_html_data(html_content)
        if scraped_data:
            save_to_csv(scraped_data, OUTPUT_CSV_FILE)
            save_to_docx(scraped_data, OUTPUT_DOCX_FILE)
        else:
            logging.warning("HTML получен, но не удалось извлечь данные. Проверьте CSS-селекторы.")
            with open("debug_page.html", "w", encoding="utf-8") as f:
                f.write(html_content)
            logging.info("HTML-код страницы сохранен в debug_page.html для анализа.")
    
    logging.info("Работа завершена.")


if __name__ == "__main__":
    main()

